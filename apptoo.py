#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import sys
import os
import re
import json
import subprocess
from datetime import datetime

from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QTextEdit, QTableWidget, QTableWidgetItem, QPushButton, QLabel,
    QLineEdit, QFileDialog, QMessageBox, QColorDialog, QDialog,
    QFormLayout, QHeaderView, QComboBox, QGroupBox, QSplitter,
    QFrame, QGraphicsDropShadowEffect, QAbstractItemView
)
from PySide6.QtCore import Qt, QDate
from PySide6.QtGui import QColor, QFontDatabase, QFont, QIcon, QPixmap, QPainter, QPainterPath, QLinearGradient

import arabic_reshaper
from bidi.algorithm import get_display

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

import openpyxl
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

CONFIG_FILE = os.path.expanduser("~/.doc_gen_config.json")
LOG_FILE = os.path.expanduser("~/.doc_gen_log.json")

# ------------------------------------------------------------------
# هوية الواجهة (ألوان احترافية موحدة تُستخدم في التنسيق العام للبرنامج)
# ------------------------------------------------------------------
PALETTE = {
    "bg":            "#F4F6F9",
    "surface":       "#FFFFFF",
    "border":        "#E1E5EB",
    "text":          "#1F2937",
    "text_muted":    "#6B7280",
    "primary":       "#2C3E82",
    "primary_dark":  "#202C63",
    "primary_light": "#EEF1FB",
    "accent":        "#3B82F6",
    "success":       "#1E8449",
    "success_dark":  "#166638",
    "danger":        "#C0392B",
    "danger_dark":   "#9E3123",
    "info":          "#2980B9",
    "info_dark":     "#20618E",
}

APP_STYLESHEET = f"""
QMainWindow, QDialog {{
    background-color: {PALETTE['bg']};
}}

QWidget {{
    color: {PALETTE['text']};
    font-family: "Segoe UI", "Cairo", "Tahoma", sans-serif;
    font-size: 13px;
}}

#headerBar {{
    background-color: {PALETTE['primary']};
    border-radius: 10px;
}}

#headerTitle {{
    color: #FFFFFF;
    font-size: 19px;
    font-weight: 700;
}}

#headerSubtitle {{
    color: #D7DCF5;
    font-size: 12px;
}}

QGroupBox {{
    background-color: {PALETTE['surface']};
    border: 1px solid {PALETTE['border']};
    border-radius: 10px;
    margin-top: 14px;
    padding: 14px 10px 10px 10px;
    font-weight: 600;
    font-size: 13px;
}}

QGroupBox::title {{
    subcontrol-origin: margin;
    subcontrol-position: top right;
    right: 12px;
    padding: 0 6px;
    color: {PALETTE['primary']};
}}

QLabel {{
    color: {PALETTE['text']};
    background: transparent;
}}

QLabel[role="fieldLabel"] {{
    color: {PALETTE['text_muted']};
    font-weight: 600;
}}

QLabel[role="footer"] {{
    color: {PALETTE['text_muted']};
    font-size: 12px;
}}

QLineEdit, QComboBox {{
    background-color: {PALETTE['surface']};
    border: 1px solid {PALETTE['border']};
    border-radius: 6px;
    padding: 6px 10px;
    min-height: 22px;
    selection-background-color: {PALETTE['accent']};
}}

QLineEdit:focus, QComboBox:focus {{
    border: 1px solid {PALETTE['accent']};
}}

QComboBox::drop-down {{
    border: none;
    width: 24px;
}}

QTextEdit {{
    background-color: {PALETTE['surface']};
    border: 1px solid {PALETTE['border']};
    border-radius: 8px;
    padding: 8px;
    selection-background-color: {PALETTE['accent']};
}}

QTextEdit:focus {{
    border: 1px solid {PALETTE['accent']};
}}

QTableWidget {{
    background-color: {PALETTE['surface']};
    alternate-background-color: {PALETTE['primary_light']};
    border: 1px solid {PALETTE['border']};
    border-radius: 8px;
    gridline-color: {PALETTE['border']};
    selection-background-color: {PALETTE['accent']};
    selection-color: #FFFFFF;
}}

QTableWidget::item {{
    padding: 4px;
}}

QHeaderView::section {{
    background-color: {PALETTE['primary']};
    color: #FFFFFF;
    font-weight: 600;
    padding: 8px;
    border: none;
    border-right: 1px solid {PALETTE['primary_dark']};
}}

QTableCornerButton::section {{
    background-color: {PALETTE['primary']};
    border: none;
}}

QSplitter::handle {{
    background-color: {PALETTE['bg']};
}}

QSplitter::handle:vertical {{
    height: 10px;
}}

QPushButton {{
    border: none;
    border-radius: 7px;
    padding: 9px 16px;
    font-weight: 600;
    color: #FFFFFF;
    background-color: {PALETTE['text_muted']};
}}

QPushButton:hover {{
    background-color: {PALETTE['primary_dark']};
}}

QPushButton#primaryAction {{
    background-color: {PALETTE['primary']};
}}
QPushButton#primaryAction:hover {{
    background-color: {PALETTE['primary_dark']};
}}

QPushButton#successAction {{
    background-color: {PALETTE['success']};
}}
QPushButton#successAction:hover {{
    background-color: {PALETTE['success_dark']};
}}

QPushButton#dangerAction {{
    background-color: {PALETTE['danger']};
}}
QPushButton#dangerAction:hover {{
    background-color: {PALETTE['danger_dark']};
}}

QPushButton#infoAction {{
    background-color: {PALETTE['info']};
}}
QPushButton#infoAction:hover {{
    background-color: {PALETTE['info_dark']};
}}

QPushButton#ghostAction {{
    background-color: transparent;
    color: {PALETTE['primary']};
    border: 1px solid {PALETTE['primary']};
}}
QPushButton#ghostAction:hover {{
    background-color: {PALETTE['primary_light']};
}}

QStatusBar {{
    background-color: {PALETTE['surface']};
    border-top: 1px solid {PALETTE['border']};
    color: {PALETTE['text_muted']};
}}

QMessageBox {{
    background-color: {PALETTE['surface']};
}}
"""


def add_card_shadow(widget, blur=22, y_offset=3, alpha=45):
    """يضيف ظلاً خفيفاً تحت العنصر لإعطائه مظهر بطاقة (card) احترافي."""
    effect = QGraphicsDropShadowEffect(widget)
    effect.setBlurRadius(blur)
    effect.setOffset(0, y_offset)
    effect.setColor(QColor(0, 0, 0, alpha))
    widget.setGraphicsEffect(effect)


def build_app_icon():
    """يبني أيقونة للتطبيق برمجياً (بدون الحاجة لملف صورة خارجي)."""
    size = 128
    pixmap = QPixmap(size, size)
    pixmap.fill(Qt.transparent)

    painter = QPainter(pixmap)
    painter.setRenderHint(QPainter.Antialiasing)

    gradient = QLinearGradient(0, 0, size, size)
    gradient.setColorAt(0, QColor(PALETTE["accent"]))
    gradient.setColorAt(1, QColor(PALETTE["primary_dark"]))

    path = QPainterPath()
    path.addRoundedRect(4, 4, size - 8, size - 8, 26, 26)
    painter.fillPath(path, gradient)

    doc_rect_w, doc_rect_h = 58, 74
    doc_x = (size - doc_rect_w) / 2
    doc_y = (size - doc_rect_h) / 2
    doc_path = QPainterPath()
    doc_path.addRoundedRect(doc_x, doc_y, doc_rect_w, doc_rect_h, 6, 6)
    painter.fillPath(doc_path, QColor("#FFFFFF"))

    painter.setPen(QColor(PALETTE["primary"]))
    line_gap = 12
    for i in range(4):
        y = doc_y + 16 + i * line_gap
        line_w = doc_rect_w - 20 if i < 3 else doc_rect_w - 34
        painter.drawLine(int(doc_x + 10), int(y), int(doc_x + 10 + line_w), int(y))

    painter.end()
    return QIcon(pixmap)


def load_config():
    default_config = {
        "last_export_dir": os.path.expanduser("~/Downloads"),
        "header_color": "#800080",
        "text_color": "#000000",
        "font_name": "Amiri",
        "font_style": "Regular"
    }
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
                default_config.update(data)
        except Exception:
            pass
    return default_config


def save_config(config_data):
    try:
        with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
            json.dump(config_data, f, ensure_ascii=False, indent=2)
    except Exception:
        pass


def add_log_entry(action, details):
    logs = []
    if os.path.exists(LOG_FILE):
        try:
            with open(LOG_FILE, 'r', encoding='utf-8') as f:
                logs = json.load(f)
        except Exception:
            logs = []
            
    entry = {
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "action": action,
        "details": details
    }
    logs.append(entry)
    
    try:
        with open(LOG_FILE, 'w', encoding='utf-8') as f:
            json.dump(logs, f, ensure_ascii=False, indent=2)
    except Exception:
        pass


def fix_text(text, max_chars=34):
    if not text:
        return ""
    text = str(text).strip()
    text = re.sub(r'[ ​‌‍‎‏﻿]', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    
    words = text.split(' ')
    lines = []
    curr_line = []
    curr_len = 0
    
    for w in words:
        word_len = len(w)
        added_len = word_len + (1 if curr_line else 0)
        if curr_len + added_len > max_chars and curr_line:
            lines.append(" ".join(curr_line))
            curr_line = [w]
            curr_len = word_len
        else:
            curr_line.append(w)
            curr_len += added_len
            
    if curr_line:
        lines.append(" ".join(curr_line))

    reshaper_config = {
        'delete_harakat': True,
        'support_ligatures': True,
        'use_unshaped_instead_of_isolated': True
    }
    reshaper = arabic_reshaper.ArabicReshaper(configuration=reshaper_config)

    processed_lines = []
    for line in lines:
        if re.search(r'[؀-ۿ]', line):
            reshaped = reshaper.reshape(line)
            reshaped = re.sub(r'[‌‍‎‏﻿]', '', reshaped)
            processed_lines.append(get_display(reshaped))
        else:
            processed_lines.append(line)

    return "<br/>".join(processed_lines)

def get_contrast_color(hex_color):
    hex_color = hex_color.lstrip('#')
    if len(hex_color) != 6:
        return "#FFFFFF"
    r, g, b = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    luminance = (0.299 * r + 0.587 * g + 0.114 * b) / 255
    return "#FFFFFF" if luminance < 0.55 else "#000000"


def find_pdf_safe_font(font_name, style_name="Regular"):
    patterns = [
        f"{font_name}:style={style_name}",
        font_name,
        f"{font_name}:lang=ar",
        "Amiri",
        ":lang=ar"
    ]
    
    for pat in patterns:
        try:
            res = subprocess.run(['fc-match', '-f', '%{file}', pat], capture_output=True, text=True)
            path = res.stdout.strip()
            if path and os.path.exists(path) and (path.endswith('.ttf') or path.endswith('.otf')):
                return path
        except Exception:
            pass

    fallback_paths = [
        "/usr/share/fonts/truetype/amiri/Amiri-Regular.ttf",
        "/usr/share/fonts/google-noto-vf/NotoSansArabic-VF.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
    ]
    for p in fallback_paths:
        if os.path.exists(p):
            return p
    return None


class SettingsDialog(QDialog):
    def __init__(self, parent=None, current_bg="#800080", current_text="#000000", current_font="Amiri", current_style="Regular"):
        super().__init__(parent)
        self.setWindowTitle("إعدادات المظهر والخطوط")
        self.resize(440, 360)
        self.setLayoutDirection(Qt.RightToLeft)
        self.setStyleSheet(APP_STYLESHEET)
        
        self.header_color = current_bg
        self.text_color = current_text
        self.font_name = current_font
        self.font_style = current_style
        
        outer = QVBoxLayout(self)
        outer.setContentsMargins(18, 18, 18, 18)
        outer.setSpacing(14)

        header = QLabel("⚙  إعدادات المظهر والخطوط")
        header.setStyleSheet(f"font-size: 15px; font-weight: 700; color: {PALETTE['primary']};")
        outer.addWidget(header)

        sub = QLabel("تحكّم بألوان الجدول والخط المستخدم في المعاينة والملفات المُصدَّرة")
        sub.setProperty("role", "footer")
        sub.setWordWrap(True)
        outer.addWidget(sub)

        line = QFrame()
        line.setFrameShape(QFrame.HLine)
        line.setStyleSheet(f"color: {PALETTE['border']};")
        outer.addWidget(line)

        form_box = QGroupBox("الألوان")
        layout = QFormLayout(form_box)
        layout.setSpacing(10)
        
        self.bg_color_btn = QPushButton("اختر لون الهيدر")
        text_contrast = get_contrast_color(self.header_color)
        self.bg_color_btn.setStyleSheet(f"background-color: {self.header_color}; color: {text_contrast}; padding: 8px; font-weight: bold; border-radius: 6px;")
        self.bg_color_btn.clicked.connect(self.choose_bg_color)
        layout.addRow("لون هيدر الجدول:", self.bg_color_btn)
        
        self.text_color_btn = QPushButton("اختر لون النص")
        self.text_color_btn.setStyleSheet(f"background-color: {self.text_color}; color: {get_contrast_color(self.text_color)}; padding: 8px; font-weight: bold; border-radius: 6px;")
        self.text_color_btn.clicked.connect(self.choose_text_color)
        layout.addRow("لون نص الجدول:", self.text_color_btn)
        outer.addWidget(form_box)
        
        font_box = QGroupBox("الخط")
        font_layout = QFormLayout(font_box)
        font_layout.setSpacing(10)

        self.font_combo = QComboBox()
        self.style_combo = QComboBox()
        
        self.load_fonts()
        self.font_combo.currentTextChanged.connect(self.on_font_changed)
        
        font_layout.addRow("نوع الخط:", self.font_combo)
        font_layout.addRow("تنسيق/نمط الخط:", self.style_combo)
        outer.addWidget(font_box)

        outer.addStretch(1)
        
        btn_box = QHBoxLayout()
        cancel_btn = QPushButton("إلغاء")
        cancel_btn.setObjectName("ghostAction")
        cancel_btn.clicked.connect(self.reject)
        btn_box.addWidget(cancel_btn)

        save_btn = QPushButton("💾 حفظ الإعدادات")
        save_btn.setObjectName("successAction")
        save_btn.clicked.connect(self.save_and_accept)
        btn_box.addWidget(save_btn)
        outer.addLayout(btn_box)
        
        self.on_font_changed(self.font_name)

    def load_fonts(self):
        families = QFontDatabase.families()
        if families:
            self.font_combo.addItems(families)
            idx = self.font_combo.findText(self.font_name)
            if idx >= 0:
                self.font_combo.setCurrentIndex(idx)
        else:
            self.font_combo.addItems(["Amiri", "DejaVu Sans", "Arial"])

    def on_font_changed(self, selected_font):
        self.style_combo.clear()
        styles = QFontDatabase.styles(selected_font)
        
        if styles and len(styles) > 1:
            self.style_combo.setEnabled(True)
            self.style_combo.addItems(styles)
            idx = self.style_combo.findText(self.font_style)
            if idx >= 0:
                self.style_combo.setCurrentIndex(idx)
        elif styles and len(styles) == 1:
            self.style_combo.addItems(styles)
            self.style_combo.setEnabled(False)
        else:
            self.style_combo.addItem("افتراضي (Regular)")
            self.style_combo.setEnabled(False)

    def choose_bg_color(self):
        col = QColorDialog.getColor(QColor(self.header_color), self)
        if col.isValid():
            self.header_color = col.name()
            text_contrast = get_contrast_color(self.header_color)
            self.bg_color_btn.setStyleSheet(f"background-color: {self.header_color}; color: {text_contrast}; padding: 8px; font-weight: bold; border-radius: 6px;")

    def choose_text_color(self):
        col = QColorDialog.getColor(QColor(self.text_color), self)
        if col.isValid():
            self.text_color = col.name()
            self.text_color_btn.setStyleSheet(f"background-color: {self.text_color}; color: {get_contrast_color(self.text_color)}; padding: 8px; font-weight: bold; border-radius: 6px;")

    def save_and_accept(self):
        self.font_name = self.font_combo.currentText()
        self.font_style = self.style_combo.currentText() if self.style_combo.isEnabled() else "Regular"
        self.accept()


class DocGeneratorApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("تطبيق توليد الكشوفات الرسمية")
        self.resize(1100, 780)
        self.setMinimumSize(880, 600)
        self.setLayoutDirection(Qt.RightToLeft)
        self.setWindowIcon(build_app_icon())
        
        self.config = load_config()
        self.header_color = self.config.get("header_color", "#800080")
        self.text_color = self.config.get("text_color", "#000000")
        self.selected_font = self.config.get("font_name", "Amiri")
        self.selected_style = self.config.get("font_style", "Regular")
        self.last_export_dir = self.config.get("last_export_dir", os.path.expanduser("~/Downloads"))
        
        self.data_rows = []
        self.init_ui()
        self.statusBar().showMessage("جاهز")
        add_log_entry("تشغيل التطبيق", "تم بدء جلسة جديدة للبرنامج.")
        
    def init_ui(self):
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        layout = QVBoxLayout(main_widget)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(12)

        # ---------- شريط العنوان العلوي ----------
        header_bar = QFrame()
        header_bar.setObjectName("headerBar")
        header_layout = QHBoxLayout(header_bar)
        header_layout.setContentsMargins(20, 14, 20, 14)

        title_box = QVBoxLayout()
        title_box.setSpacing(2)
        title_lbl = QLabel("📋  نظام توليد الكشوفات الرسمية")
        title_lbl.setObjectName("headerTitle")
        subtitle_lbl = QLabel("معالجة الأسماء وبيانات الجوازات وتصديرها بصيغ PDF وExcel وWord")
        subtitle_lbl.setObjectName("headerSubtitle")
        title_box.addWidget(title_lbl)
        title_box.addWidget(subtitle_lbl)
        header_layout.addLayout(title_box)
        header_layout.addStretch(1)

        add_card_shadow(header_bar, blur=24, y_offset=4, alpha=60)
        layout.addWidget(header_bar)

        # ---------- لوحة بيانات الكشف والتحكم ----------
        inputs_group = QGroupBox("بيانات الكشف والتحكم")
        inputs_layout = QHBoxLayout(inputs_group)
        inputs_layout.setSpacing(10)
        
        batch_lbl = QLabel("رقم الموافقة:")
        batch_lbl.setProperty("role", "fieldLabel")
        inputs_layout.addWidget(batch_lbl)
        self.batch_input = QLineEdit("8895")
        inputs_layout.addWidget(self.batch_input)
        
        date_lbl = QLabel("تاريخ الإصدار:")
        date_lbl.setProperty("role", "fieldLabel")
        inputs_layout.addWidget(date_lbl)
        self.date_input = QLineEdit(QDate.currentDate().toString("dd/MM/yyyy"))
        inputs_layout.addWidget(self.date_input)
        
        inputs_layout.addStretch(1)

        settings_btn = QPushButton("⚙  الإعدادات والألوان")
        settings_btn.setObjectName("ghostAction")
        settings_btn.clicked.connect(self.open_settings)
        inputs_layout.addWidget(settings_btn)
        
        add_card_shadow(inputs_group)
        layout.addWidget(inputs_group)
        
        splitter = QSplitter(Qt.Vertical)
        
        text_container = QWidget()
        text_layout = QVBoxLayout(text_container)
        text_layout.setContentsMargins(0, 0, 0, 0)
        text_layout.setSpacing(6)
        text_hdr_lbl = QLabel("✏  ألصق النص هنا:")
        text_hdr_lbl.setProperty("role", "fieldLabel")
        text_layout.addWidget(text_hdr_lbl)
        self.text_edit = QTextEdit()
        self.text_edit.setPlaceholderText("ألصق النص المحتوي على الأسماء، الجوازات، وتواريخ الميلاد...")
        add_card_shadow(self.text_edit, blur=16, alpha=30)
        text_layout.addWidget(self.text_edit)
        splitter.addWidget(text_container)
        
        table_container = QWidget()
        table_layout = QVBoxLayout(table_container)
        table_layout.setContentsMargins(0, 0, 0, 0)
        table_layout.setSpacing(6)

        table_hdr_layout = QHBoxLayout()
        table_hdr_lbl = QLabel("📑  معاينة البيانات المعالجة:")
        table_hdr_lbl.setProperty("role", "fieldLabel")
        table_hdr_layout.addWidget(table_hdr_lbl)
        table_hdr_layout.addStretch(1)
        self.row_count_lbl = QLabel("عدد السجلات: 0")
        self.row_count_lbl.setProperty("role", "footer")
        table_hdr_layout.addWidget(self.row_count_lbl)
        table_layout.addLayout(table_hdr_layout)
        
        self.table = QTableWidget()
        self.table.setColumnCount(6)
        self.table.setHorizontalHeaderLabels(["ت", "الاسم", "تاريخ الميلاد", "رقم الجواز", "تاريخ الإصدار", "رقم الموافقة"])
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.table.verticalHeader().setVisible(False)
        self.table.setAlternatingRowColors(True)
        self.table.setSelectionBehavior(QAbstractItemView.SelectRows)
        add_card_shadow(self.table, blur=16, alpha=30)
        table_layout.addWidget(self.table)
        splitter.addWidget(table_container)
        splitter.setSizes([260, 420])
        
        layout.addWidget(splitter, 1)
        
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(10)
        
        process_btn = QPushButton("🔍 معالجة وفحص البيانات")
        process_btn.setObjectName("successAction")
        process_btn.clicked.connect(self.process_data)
        btn_layout.addWidget(process_btn)

        btn_layout.addStretch(1)
        
        export_pdf_btn = QPushButton("📄 تصدير PDF")
        export_pdf_btn.setObjectName("dangerAction")
        export_pdf_btn.clicked.connect(lambda: self.export_data('pdf'))
        btn_layout.addWidget(export_pdf_btn)
        
        export_excel_btn = QPushButton("📊 تصدير Excel")
        export_excel_btn.setObjectName("primaryAction")
        export_excel_btn.clicked.connect(lambda: self.export_data('excel'))
        btn_layout.addWidget(export_excel_btn)
        
        export_word_btn = QPushButton("📝 تصدير Word")
        export_word_btn.setObjectName("infoAction")
        export_word_btn.clicked.connect(lambda: self.export_data('word'))
        btn_layout.addWidget(export_word_btn)
        
        layout.addLayout(btn_layout)

    def open_settings(self):
        dlg = SettingsDialog(self, self.header_color, self.text_color, self.selected_font, self.selected_style)
        if dlg.exec():
            self.header_color = dlg.header_color
            self.text_color = dlg.text_color
            self.selected_font = dlg.font_name
            self.selected_style = dlg.font_style
            
            self.config["header_color"] = self.header_color
            self.config["text_color"] = self.text_color
            self.config["font_name"] = self.selected_font
            self.config["font_style"] = self.selected_style
            save_config(self.config)
            
            add_log_entry("تحديث الإعدادات", f"الخط: {self.selected_font}, التنسيق: {self.selected_style}, لون الهيدر: {self.header_color}")
            self.statusBar().showMessage("تم حفظ إعدادات الألوان والخطوط", 5000)
            QMessageBox.information(self, "نجاح", "تم حفظ إعدادات الألوان والخطوط بنجاح.")

    def parse_date(self, text):
        if not text:
            return ""
        text = str(text).strip()
        digits = re.sub(r'[^\d]', '', text)
        
        if len(digits) == 8:
            if digits.startswith(('19', '20')):
                yyyy = digits[:4]
                mm = digits[4:6]
                dd = digits[6:8]
                return f"{dd}/{mm}/{yyyy}"
            else:
                dd = digits[:2]
                mm = digits[2:4]
                yyyy = digits[4:8]
                return f"{dd}/{mm}/{yyyy}"
                
        parts = [p for p in re.split(r'[^\d]', text) if p]
        if len(parts) == 3:
            p1, p2, p3 = parts[0], parts[1], parts[2]
            if len(p1) == 4:
                return f"{p3.zfill(2)}/{p2.zfill(2)}/{p1}"
            elif len(p3) == 4:
                return f"{p1.zfill(2)}/{p2.zfill(2)}/{p3}"
                
        return text

    def process_data(self):
        raw_text = self.text_edit.toPlainText().strip()
        if not raw_text:
            QMessageBox.warning(self, "تنبيه", "يرجى إلصاق النص أولاً!")
            return
            
        lines = [line.strip() for line in raw_text.split('\n') if line.strip() and not line.strip().startswith('<<<')]
        
        names, passports, dobs = [], [], []
        passport_pattern = re.compile(r'^[A-Za-z0-9]{6,}$')
        date_pattern = re.compile(r'^(\d{8}|\d{1,4}[-/.]\d{1,2}[-/.]\d{1,4})$')
        
        for line in lines:
            if date_pattern.match(line):
                dobs.append(self.parse_date(line))
            elif passport_pattern.match(line) and any(c.isdigit() for c in line) and not re.search(r'[\u0600-\u06FF]', line):
                passports.append(line)
            else:
                names.append(line)
                
        if not (len(names) == len(passports) == len(dobs)):
            msg = f"أعداد البيانات غير متطابقة! أسماء: {len(names)}، جوازات: {len(passports)}، مواليد: {len(dobs)}"
            add_log_entry("خطأ معالجة", msg)
            QMessageBox.critical(self, "خطأ في المطابقة", msg)
            return
            
        self.data_rows = []
        batch_no = self.batch_input.text().strip()
        issue_date = self.parse_date(self.date_input.text().strip())
        
        for i in range(len(names)):
            self.data_rows.append({
                "seq": i + 1,
                "name": names[i],
                "dob": dobs[i],
                "passport": passports[i],
                "issue_date": issue_date,
                "batch": batch_no
            })
            
        self.display_table()
        add_log_entry("معالجة بيانات ناجحة", f"تمت معالجة {len(self.data_rows)} سجل للوجبة {batch_no}.")
        self.statusBar().showMessage(f"تم معالجة {len(self.data_rows)} سجل بنجاح", 5000)
        QMessageBox.information(self, "نجاح", f"تم معالجة {len(self.data_rows)} سجل بنجاح!")

    def display_table(self):
        self.table.setRowCount(len(self.data_rows))
        for r, row in enumerate(self.data_rows):
            self.table.setItem(r, 0, QTableWidgetItem(str(row["seq"])))
            self.table.setItem(r, 1, QTableWidgetItem(row["name"]))
            self.table.setItem(r, 2, QTableWidgetItem(row["dob"]))
            self.table.setItem(r, 3, QTableWidgetItem(row["passport"]))
            self.table.setItem(r, 4, QTableWidgetItem(row["issue_date"]))
            self.table.setItem(r, 5, QTableWidgetItem(row["batch"]))
        self.row_count_lbl.setText(f"عدد السجلات: {len(self.data_rows)}")

    def get_auto_filename(self, fmt):
        ext_map = {'excel': 'xlsx', 'word': 'docx', 'pdf': 'pdf'}
        ext = ext_map.get(fmt.lower(), fmt)
        
        today_date = datetime.now().strftime('%d-%m-%Y')
        base_name = f"كشف_موافقات_{today_date}"
        target_dir = self.last_export_dir if os.path.exists(self.last_export_dir) else os.path.expanduser("~/Downloads")
        
        counter = 0
        while True:
            suffix = "" if counter == 0 else f"_{counter}"
            candidate_name = f"{base_name}{suffix}.{ext}"
            full_path = os.path.join(target_dir, candidate_name)
            if not os.path.exists(full_path):
                break
            counter += 1
            
        filter_str = f"{ext.upper()} Files (*.{ext})"
        path, _ = QFileDialog.getSaveFileName(
            self, "حفظ الملف", os.path.join(target_dir, candidate_name),
            filter_str
        )
        
        if path:
            if not path.endswith(f".{ext}"):
                path += f".{ext}"
            self.last_export_dir = os.path.dirname(path)
            self.config["last_export_dir"] = self.last_export_dir
            save_config(self.config)
            
        return path

    def export_data(self, fmt):
        if not self.data_rows:
            QMessageBox.warning(self, "تنبيه", "يرجى معالجة البيانات أولاً قبل التصدير!")
            return
            
        path = self.get_auto_filename(fmt)
        if not path:
            return
            
        try:
            if fmt == 'pdf':
                self.export_pdf(path)
            elif fmt == 'excel':
                self.export_excel(path)
            elif fmt == 'word':
                self.export_word(path)
                
            add_log_entry(f"تصدير {fmt.upper()}", f"تم تصدير {len(self.data_rows)} سجل إلى المسار: {path}")
            self.statusBar().showMessage(f"تم حفظ الملف: {path}", 6000)
            QMessageBox.information(self, "نجاح التصدير", f"تم حفظ الملف بنجاح في:\n{path}")
        except Exception as e:
            add_log_entry(f"خطأ تصدير {fmt.upper()}", str(e))
            self.statusBar().showMessage("حدث خطأ أثناء التصدير", 6000)
            QMessageBox.critical(self, "خطأ أثناء التصدير", f"حدث خطأ أثناء حفظ الملف:\n{str(e)}")

    def export_pdf(self, path):
        font_path = find_pdf_safe_font(self.selected_font, self.selected_style)
        
        if font_path and os.path.exists(font_path):
            font_alias = f"Font_{abs(hash(font_path))}"
            try:
                pdfmetrics.registerFont(TTFont(font_alias, font_path))
                add_log_entry("تطبيق خط PDF", f"تم تحميل ملف الخط من: {font_path}")
            except Exception as e:
                font_alias = "Helvetica"
                add_log_entry("خطأ خط PDF", f"تعذر تسجيل الخط: {str(e)}")
        else:
            font_alias = "Helvetica"

        doc = SimpleDocTemplate(
            path, pagesize=A4, rightMargin=20, leftMargin=20, topMargin=20, bottomMargin=20
        )
        elements = []
        
        header_text_color = get_contrast_color(self.header_color)
        body_text_color = self.text_color

        normal_style = ParagraphStyle(
            'ArabicStyle',
            fontName=font_alias,
            fontSize=10,
            leading=13,
            alignment=1
        )

        headers = ["رقم الموافقة", "تاريخ الإصدار", "رقم الجواز", "تاريخ الميلاد", "الاسم", "ت"]
        reshaped_headers = [
            Paragraph(f'<font color="{header_text_color}">{fix_text(h)}</font>', normal_style) 
            for h in headers
        ]
        
        table_data = [reshaped_headers]
        
        for row in self.data_rows:
            table_data.append([
                Paragraph(f'<font color="{body_text_color}">{row["batch"]}</font>', normal_style),
                Paragraph(f'<font color="{body_text_color}">{row["issue_date"]}</font>', normal_style),
                Paragraph(f'<font color="{body_text_color}">{row["passport"]}</font>', normal_style),
                Paragraph(f'<font color="{body_text_color}">{row["dob"]}</font>', normal_style),
                Paragraph(f'<font color="{body_text_color}">{fix_text(row["name"])}</font>', normal_style),
                Paragraph(f'<font color="{body_text_color}">{row["seq"]}</font>', normal_style)
            ])
            
        col_widths = [80, 80, 90, 90, 180, 30]
        hex_bg = self.header_color.lstrip('#')
        rgb_bg = colors.HexColor(f"#{hex_bg}")
        
        t = Table(table_data, colWidths=col_widths, repeatRows=1)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), rgb_bg),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('GRID', (0, 0), (-1, -1), 0.8, colors.black),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
        ]))
        
        elements.append(t)
        doc.build(elements)

    def export_excel(self, path):
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "الكشف"
        ws.views.sheetView[0].rightToLeft = True
        
        headers = ["ت", "الاسم", "تاريخ الميلاد", "رقم الجواز", "تاريخ الإصدار", "رقم الموافقة"]
        ws.append(headers)
        
        hex_bg = self.header_color.lstrip('#')
        hex_hdr_txt = get_contrast_color(self.header_color).lstrip('#')
        hex_body_txt = self.text_color.lstrip('#')
        
        fill = PatternFill(start_color=hex_bg, end_color=hex_bg, fill_type="solid")
        is_bold_style = "bold" in self.selected_style.lower() or "heavy" in self.selected_style.lower()
        
        header_font = Font(name=self.selected_font, size=11, bold=True, color=hex_hdr_txt)
        data_font = Font(name=self.selected_font, size=11, bold=is_bold_style, color=hex_body_txt)
        
        thin_border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin')
        )
        
        for col_num in range(1, 7):
            cell = ws.cell(row=1, column=col_num)
            cell.fill = fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = thin_border
            
        for row in self.data_rows:
            r_data = [row["seq"], row["name"], row["dob"], row["passport"], row["issue_date"], row["batch"]]
            ws.append(r_data)
            
        for row in ws.iter_rows(min_row=2, max_row=len(self.data_rows)+1, min_col=1, max_col=6):
            for idx, cell in enumerate(row):
                cell.font = data_font
                cell.border = thin_border
                align = "right" if idx == 1 else "center"
                cell.alignment = Alignment(horizontal=align, vertical="center")
                
        for col in ws.columns:
            max_len = max(len(str(cell.value or '')) for cell in col)
            col_letter = openpyxl.utils.get_column_letter(col[0].column)
            ws.column_dimensions[col_letter].width = max(max_len + 6, 16)
            
        wb.save(path)
        try:
            os.chmod(path, 0o666)
        except Exception:
            pass

    def export_word(self, path):
        doc = docx.Document()
        
        # ضبط هوامش الصفحة لتستوعب الجدول بأبعاد متناسقة
        for section in doc.sections:
            section.top_margin = docx.shared.Cm(1.5)
            section.bottom_margin = docx.shared.Cm(1.5)
            section.left_margin = docx.shared.Cm(1.5)
            section.right_margin = docx.shared.Cm(1.5)
            
        table = doc.add_table(rows=1, cols=6)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        
        tblPr = table._tbl.tblPr
        tblCellMar = parse_xml(f'<w:bidiVisual {nsdecls("w")}/>')
        tblPr.append(tblCellMar)
        
        # تحديد أعراض ثابتة ومناسبة لكل عمود لمنع انقسام التواريخ والأرقام
        col_widths = [
            docx.shared.Cm(1.0),   # ت
            docx.shared.Cm(6.0),   # الاسم
            docx.shared.Cm(2.8),   # تاريخ الميلاد
            docx.shared.Cm(2.8),   # رقم الجواز
            docx.shared.Cm(2.8),   # تاريخ الإصدار
            docx.shared.Cm(2.6)    # رقم الموافقة
        ]
        
        headers = ["ت", "الاسم", "تاريخ الميلاد", "رقم الجواز", "تاريخ الإصدار", "رقم الموافقة"]
        hdr_cells = table.rows[0].cells
        hex_bg = self.header_color.lstrip('#')
        
        for i, header_text in enumerate(headers):
            hdr_cells[i].text = header_text
            hdr_cells[i].width = col_widths[i]
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                run = p.runs[0]
                run.font.name = self.selected_font
                run.font.bold = True
                rPr = run._r.get_or_add_rPr()
                rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{self.selected_font}" w:hAnsi="{self.selected_font}" w:cs="{self.selected_font}"/>')
                rPr.append(rFonts)
                
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_bg}"/>')
            hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
            
        hex_txt = self.text_color.lstrip('#')
        r, g, b = tuple(int(hex_txt[i:i+2], 16) for i in (0, 2, 4))
        word_rgb_color = RGBColor(r, g, b)
        is_bold_style = "bold" in self.selected_style.lower() or "heavy" in self.selected_style.lower()

        for row in self.data_rows:
            row_cells = table.add_row().cells
            row_data = [str(row["seq"]), row["name"], row["dob"], row["passport"], row["issue_date"], row["batch"]]
            for i, val in enumerate(row_data):
                row_cells[i].text = val
                row_cells[i].width = col_widths[i]
                p = row_cells[i].paragraphs[0]
                align = WD_ALIGN_PARAGRAPH.RIGHT if i == 1 else WD_ALIGN_PARAGRAPH.CENTER
                p.alignment = align
                if p.runs:
                    run = p.runs[0]
                    run.font.name = self.selected_font
                    run.font.bold = is_bold_style
                    run.font.color.rgb = word_rgb_color
                    rPr = run._r.get_or_add_rPr()
                    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{self.selected_font}" w:hAnsi="{self.selected_font}" w:cs="{self.selected_font}"/>')
                    rPr.append(rFonts)
                
        doc.save(path)
        try:
            os.chmod(path, 0o666)
        except Exception:
            pass


if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setStyleSheet(APP_STYLESHEET)
    app.setFont(QFont("Segoe UI", 10))
    window = DocGeneratorApp()
    window.show()
    sys.exit(app.exec())
